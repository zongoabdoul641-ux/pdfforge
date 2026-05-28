api/convert.py — Vercel Serverless Function
Handles PDF conversion requests.
"""

import json
import base64
import io
import os
import tempfile
from http.server import BaseHTTPRequestHandler


# ─── Conversion dispatcher ───────────────────────────────────────────────────

def convert_to_pdf(file_bytes: bytes, filename: str, quality: str) -> bytes:
    """Convert various file types to PDF."""
    ext = filename.rsplit(".", 1)[-1].lower()

    if ext in ("jpg", "jpeg", "png", "gif", "bmp", "webp"):
        return image_to_pdf(file_bytes, quality)
    elif ext in ("docx", "doc"):
        return docx_to_pdf(file_bytes, quality)
    elif ext in ("xlsx", "xls"):
        return xlsx_to_pdf(file_bytes, quality)
    elif ext == "txt":
        return txt_to_pdf(file_bytes, quality)
    elif ext == "pdf":
        return optimize_pdf(file_bytes, quality)
    else:
        raise ValueError(f"Format non supporté : .{ext}")


# ─── Converters ──────────────────────────────────────────────────────────────

def image_to_pdf(file_bytes: bytes, quality: str) -> bytes:
    """Convert image to PDF using ReportLab + Pillow."""
    from PIL import Image
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    img = Image.open(io.BytesIO(file_bytes))

    # Convert RGBA/palette to RGB
    if img.mode in ("RGBA", "P", "LA"):
        background = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode == "P":
            img = img.convert("RGBA")
        if "A" in img.mode:
            background.paste(img, mask=img.split()[-1])
        else:
            background.paste(img)
        img = background
    elif img.mode != "RGB":
        img = img.convert("RGB")

    # Quality settings
    dpi_map = {"rapide": 72, "standard": 150, "haute": 300}
    dpi = dpi_map.get(quality, 150)

    # Resize if too large (~30 inches max)
    max_px = dpi * 30
    if max(img.size) > max_px:
        ratio = max_px / max(img.size)
        img = img.resize(
            (int(img.width * ratio), int(img.height * ratio)),
            Image.LANCZOS,
        )

    # Save temp image
    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp_img:
        img.save(tmp_img.name, "JPEG", quality=85 if quality != "haute" else 95)
        tmp_img_path = tmp_img.name

    # Build PDF
    pdf_buffer = io.BytesIO()
    page_w, page_h = A4

    img_w, img_h = img.size
    ratio = min(page_w / img_w, page_h / img_h) * 0.9
    draw_w, draw_h = img_w * ratio, img_h * ratio
    x = (page_w - draw_w) / 2
    y = (page_h - draw_h) / 2

    c = canvas.Canvas(pdf_buffer, pagesize=A4)
    c.drawImage(tmp_img_path, x, y, width=draw_w, height=draw_h)
    c.save()
    os.unlink(tmp_img_path)

    return pdf_buffer.getvalue()


def docx_to_pdf(file_bytes: bytes, quality: str) -> bytes:
    """Convert DOCX to PDF using python-docx + ReportLab."""
    from docx import Document
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    doc = Document(io.BytesIO(file_bytes))
    pdf_buffer = io.BytesIO()

    pdf_doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading1"],
        fontSize=16,
        spaceAfter=10,
        textColor=colors.HexColor("#0d0d0d"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=6,
    )

    story = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue

        style_name = para.style.name or ""
        if "Heading 1" in style_name or "Title" in style_name:
            story.append(Paragraph(text, heading_style))
        elif "Heading" in style_name:
            story.append(Paragraph(text, styles["Heading2"]))
        else:
            story.append(Paragraph(text, body_style))

    # Tables
    for table in doc.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        if data:
            t = Table(data)
            t.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1,  0), colors.HexColor("#f5f0e8")),
                ("TEXTCOLOR",    (0, 0), (-1,  0), colors.HexColor("#0d0d0d")),
                ("FONTNAME",     (0, 0), (-1,  0), "Helvetica-Bold"),
                ("FONTSIZE",     (0, 0), (-1, -1), 9),
                ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#d9d1c4")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#faf7f2")]),
                ("PADDING",      (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

    if not story:
        story.append(Paragraph("Document vide", body_style))

    pdf_doc.build(story)
    return pdf_buffer.getvalue()


def xlsx_to_pdf(file_bytes: bytes, quality: str) -> bytes:
    """Convert XLSX to PDF using openpyxl + ReportLab."""
    import openpyxl
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
    )

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    pdf_buffer = io.BytesIO()

    pdf_doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=landscape(A4),
        rightMargin=1.5 * cm,
        leftMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    styles = getSampleStyleSheet()
    story = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        story.append(Paragraph(f"Feuille : {sheet_name}", styles["Heading2"]))
        story.append(Spacer(1, 8))

        data = []
        for row in ws.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            if any(row_data):
                data.append(row_data)

        if data:
            col_count = max(len(r) for r in data)
            data = [r + [""] * (col_count - len(r)) for r in data]  # normalize row lengths

            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1,  0), colors.HexColor("#c94f1e")),
                ("TEXTCOLOR",    (0, 0), (-1,  0), colors.white),
                ("FONTNAME",     (0, 0), (-1,  0), "Helvetica-Bold"),
                ("FONTSIZE",     (0, 0), (-1, -1), 8),
                ("GRID",         (0, 0), (-1, -1), 0.4, colors.HexColor("#d9d1c4")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#faf7f2")]),
                ("PADDING",      (0, 0), (-1, -1), 5),
                ("VALIGN",       (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(t)

        story.append(Spacer(1, 20))

    pdf_doc.build(story)
    return pdf_buffer.getvalue()


def txt_to_pdf(file_bytes: bytes, quality: str) -> bytes:
    """Convert plain text to PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    text = file_bytes.decode("utf-8", errors="replace")
    pdf_buffer = io.BytesIO()

    pdf_doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    mono_style = ParagraphStyle(
        "Mono",
        parent=styles["Normal"],
        fontName="Courier",
        fontSize=10,
        leading=14,
    )

    story = []
    for line in text.splitlines():
        if line.strip():
            story.append(
                Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), mono_style)
            )
        else:
            story.append(Spacer(1, 6))

    pdf_doc.build(story)
    return pdf_buffer.getvalue()


def optimize_pdf(file_bytes: bytes, quality: str) -> bytes:
    """Optimize/compress an existing PDF."""
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(io.BytesIO(file_bytes))
    writer = PdfWriter()

    for page in reader.pages:
        if quality == "rapide":
            page.compress_content_streams()
        writer.add_page(page)

    writer.add_metadata(reader.metadata or {})

    if quality in ("rapide", "standard"):
        writer.compress_identical_objects(remove_identicals=True, remove_orphans=True)

    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


# ─── Vercel handler ──────────────────────────────────────────────────────────

class handler(BaseHTTPRequestHandler):

    def do_OPTIONS(self):
        self.send_response(200)
        self._cors()
        self.end_headers()

    def do_POST(self):
        try:
            length    = int(self.headers.get("Content-Length", 0))
            body      = json.loads(self.rfile.read(length))
            filename  = body.get("filename", "file.pdf")
            quality   = body.get("quality", "standard")
            file_b64  = body.get("file")

            if not file_b64:
                self._error(400, "Aucun fichier reçu.")
                return

            file_bytes = base64.b64decode(file_b64)
            pdf_bytes  = convert_to_pdf(file_bytes, filename, quality)

            result = {
                "success":  True,
                "filename": os.path.splitext(filename)[0] + ".pdf",
                "size":     len(pdf_bytes),
                "data":     base64.b64encode(pdf_bytes).decode(),
            }

            self.send_response(200)
            self._cors()
            self.send_header("Content-Type", "application/json")
            self.end_headers()
            self.wfile.write(json.dumps(result).encode())

        except ValueError as e:
            self._error(400, str(e))
        except Exception as e:
            self._error(500, f"Erreur serveur : {str(e)}")

    def _cors(self):
        self.send_header("Access-Control-Allow-Origin",  "*")
        self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")

    def _error(self, code: int, msg: str):
        self.send_response(code)
        self._cors()
        self.send_header("Content-Type", "application/json")
        self.end_headers()
        self.wfile.write(json.dumps({"success": False, "error": msg}).encode())

    def log_message(self, *args):
        pass  # Silence logs in Vercel
